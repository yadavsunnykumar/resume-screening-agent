import io
import pdfplumber
from docx import Document


def extract_text(filename: str, content: bytes) -> str:
    """
    Extract plain text from an uploaded file.
    Dispatches by file extension to the appropriate parser.
    Returns a plain text string.
    """
    extension = filename.rsplit(".", 1)[-1].lower()

    if extension == "pdf":
        return _extract_pdf(content)
    elif extension == "docx":
        return _extract_docx(content)
    elif extension == "txt":
        return _extract_txt(content)
    else:
        raise ValueError(f"Unsupported file type: .{extension}. Please upload PDF, DOCX, or TXT.")


def _extract_pdf(content: bytes) -> str:
    """Extract text from a PDF file using pdfplumber."""
    text_parts = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n\n".join(text_parts)


def _extract_docx(content: bytes) -> str:
    """
    Extract text from a DOCX file.
    Handles both paragraphs and table cells (for multi-column resume templates).
    Deduplicates merged cells using id() tracking.
    """
    document = Document(io.BytesIO(content))
    text_parts = []

    # Extract paragraphs
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Extract table cells (handles multi-column resume layouts)
    seen = set()
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if id(cell) in seen:
                    continue
                seen.add(id(cell))
                if cell.text.strip():
                    text_parts.append(cell.text)

    return "\n".join(text_parts)


def _extract_txt(content: bytes) -> str:
    """Extract text from a plain text file."""
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("utf-8", errors="ignore")
